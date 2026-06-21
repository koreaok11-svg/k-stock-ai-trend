# -*- coding: utf-8 -*-
"""
신의 계시
GitHub + Render 배포용 Streamlit 단일 파일

파일명: app.py

requirements.txt:
streamlit
pandas
openai
python-docx
pypdf

Render Start Command:
streamlit run app.py --server.port $PORT --server.address 0.0.0.0

환경변수:
OPENAI_API_KEY=OpenAI API Key
OPENAI_MODEL=gpt-5.5
OPENAI_INITIAL_CREDIT=10
OPENAI_INPUT_COST_PER_1M=1.25
OPENAI_OUTPUT_COST_PER_1M=10.00

V37 핵심:
- V31 구조를 유지하면서 작품보관함 편집/삭제/다운로드 기능 강화
- 메뉴: 내 작품 / 작품 생성 / 작품 수정 / 작품 보관함 / 다운로드 / OpenAI / AI 공동작가 / 가이드
- AI 생성/수정/작품보관함 선택버전 이어쓰기/AI 지시 실행 성공 시 자동저장 + 자동 버전 생성
- 수정 결과가 앱을 나갔다 들어와도 사라지지 않도록 프로젝트 JSON에 저장
- 작품 보관함에서 버전별 확인/편집/복원/다운로드/삭제 가능, V1/V2 보호 유지
- 다중 작품 관리 준비: projects_v30 폴더에 작품별 JSON 저장
"""

import os
import re
import io
import json
from pathlib import Path
from datetime import datetime

import pandas as pd
import streamlit as st

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

APP_VERSION = "SHIN_OF_REVELATION_V37"
APP_TITLE = "신의 계시"

# V37에서도 기존 V30/V31/V2 작업물이 사라지지 않도록 같은 저장 폴더를 사용합니다.
PROJECT_DIR = Path("projects_v30")
PROJECT_DIR.mkdir(exist_ok=True)
INDEX_FILE = PROJECT_DIR / "_index.json"

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini").strip()
OPENAI_INITIAL_CREDIT = float(os.getenv("OPENAI_INITIAL_CREDIT", "10") or 10)
OPENAI_INPUT_COST_PER_1M = float(os.getenv("OPENAI_INPUT_COST_PER_1M", "1.25") or 1.25)
OPENAI_OUTPUT_COST_PER_1M = float(os.getenv("OPENAI_OUTPUT_COST_PER_1M", "10.00") or 10.00)

st.set_page_config(page_title=APP_TITLE, page_icon="✍️", layout="wide")

st.markdown("""
<style>
.stApp {background: radial-gradient(circle at top left, rgba(216,243,220,.85), transparent 34%), radial-gradient(circle at top right, rgba(169,222,249,.65), transparent 28%), linear-gradient(180deg, #fffaf0 0%, #f8fff4 100%);} 
.block-container {max-width:1180px; padding-top:3.2rem !important;}
.main-title {
  font-size:clamp(28px, 4.2vw, 44px);
  font-weight:950;
  color:#4b3b2f;
  line-height:1.45;
  word-break:keep-all;
  margin-top:18px;
  padding-top:18px;
  padding-bottom:8px;
  overflow:visible;
  white-space:normal;
}
.sub-title {font-size:14px; color:#6d6259; margin-bottom:16px; word-break:keep-all;}
.card {border:1px solid rgba(121,183,145,.28); border-radius:22px; padding:18px; background:rgba(255,255,255,.92); box-shadow:0 8px 22px rgba(91,70,54,.08); margin-bottom:14px;}
.ai-box {border:1px solid rgba(121,183,145,.32); border-radius:20px; padding:16px 18px; background:linear-gradient(135deg, rgba(255,243,176,.75), rgba(216,243,220,.85)); line-height:1.75; white-space:pre-wrap; word-break:keep-all;}
.memo-box {border:1px solid rgba(169,222,249,.45); border-radius:18px; padding:14px 16px; background:rgba(255,255,255,.78); line-height:1.65; white-space:pre-wrap; word-break:keep-all;}
.notice {border-radius:18px; padding:14px 16px; background:rgba(216,243,220,.75); border:1px solid rgba(121,183,145,.35); margin-bottom:12px;}
.stTabs [data-baseweb="tab-list"] {gap:8px; flex-wrap:wrap;}
.stTabs [data-baseweb="tab"] {background:rgba(255,255,255,.78); border-radius:999px; padding:8px 12px; border:1px solid rgba(121,183,145,.25);} 
.stTabs [aria-selected="true"] {background:#d8f3dc !important; color:#2d6a4f !important;}
.stButton>button, .stDownloadButton>button, a[data-testid="stLinkButton"] {border-radius:999px !important; border:1px solid rgba(121,183,145,.45) !important; background:linear-gradient(135deg, #fff3b0, #d8f3dc) !important; color:#4b3b2f !important; font-weight:800 !important;}
div[data-testid="stMetric"] {background:rgba(255,255,255,.92); border:1px solid rgba(121,183,145,.25); border-radius:18px; padding:14px; box-shadow:0 6px 18px rgba(91,70,54,.08); min-height:104px;}
textarea {font-size:15px !important;}
@media (max-width:768px){.block-container{padding:.8rem .7rem}.main-title{font-size:25px!important}.sub-title{font-size:12px}.card{padding:14px}.ai-box{font-size:14px}.memo-box{font-size:14px}}

/* V37 가독성 개선 */
html, body, [class*="css"] {font-family: "Apple SD Gothic Neo", "Noto Sans KR", "Malgun Gothic", sans-serif;}
.block-container p, .stMarkdown, .stText, label {line-height:1.85 !important; letter-spacing:-0.01em;}
.ai-box, .memo-box, textarea {font-size:17px !important; line-height:1.95 !important;}
.stDataFrame, div[data-testid="stTable"] {font-size:15px !important;}
.card h1, .card h2, .card h3 {line-height:1.45 !important;}
.version-toolbox {border:1px solid rgba(121,183,145,.35); border-radius:20px; padding:16px; background:rgba(255,255,255,.86); margin:12px 0;}
.danger-box {border:1px solid rgba(214,40,40,.35); border-radius:18px; padding:14px; background:rgba(255,235,235,.85);}
@media (max-width:768px){
  .ai-box, .memo-box, textarea {font-size:18px !important; line-height:2.05 !important;}
  .stButton>button, .stDownloadButton>button {font-size:16px !important; padding:.65rem .8rem !important;}
}


/* V37 모바일/보관함 압축 가독성 개선 */
@media (max-width:768px){
  .block-container {padding-top:1.4rem !important; padding-left:.65rem !important; padding-right:.65rem !important;}
  .main-title {font-size:22px !important; line-height:1.25 !important; margin-top:4px !important; padding-top:8px !important; padding-bottom:2px !important;}
  .sub-title {font-size:11px !important; line-height:1.35 !important; margin-bottom:8px !important;}
  div[data-testid="stMetric"] {min-height:78px !important; padding:10px !important;}
  div[data-testid="stMetricValue"] {font-size:22px !important;}
  .card {padding:12px !important; border-radius:16px !important; margin-bottom:8px !important; font-size:14px !important; line-height:1.45 !important;}
  .notice {padding:10px 12px !important; border-radius:14px !important; font-size:13px !important; line-height:1.45 !important;}
  .memo-box, .ai-box {font-size:14px !important; line-height:1.55 !important; padding:12px !important; border-radius:14px !important;}
  .memo-box.compact-preview {max-height:360px; overflow-y:auto;}
  textarea {font-size:14px !important; line-height:1.55 !important;}
  .stTabs [data-baseweb="tab"] {padding:6px 9px !important; font-size:13px !important;}
}
.memo-box.compact-preview {
  max-height:480px;
  overflow-y:auto;
  font-size:15px !important;
  line-height:1.6 !important;
}


/* V37 상단 잘림 방지 + 최신버전 고정 표시 */
.block-container {padding-top:4.2rem !important;}
.top-version-card {
  margin-top:10px;
  margin-bottom:12px;
  padding:14px 16px 12px 16px;
  border-radius:22px;
  background:rgba(255,255,255,.74);
  border:1px solid rgba(121,183,145,.22);
  box-shadow:0 8px 20px rgba(91,70,54,.06);
}
.top-version-badge {
  display:inline-block;
  padding:5px 11px;
  border-radius:999px;
  background:#d8f3dc;
  color:#1b6b4a;
  font-size:13px;
  font-weight:900;
  margin-bottom:6px;
}
.main-title {
  margin-top:0 !important;
  padding-top:0 !important;
  font-size:clamp(24px, 3.5vw, 40px) !important;
  line-height:1.25 !important;
}
.sub-title {
  line-height:1.35 !important;
}
@media (max-width:768px){
  .block-container {padding-top:5.8rem !important;}
  .top-version-card {
    margin-top:8px !important;
    padding:12px 12px 10px 12px !important;
    border-radius:18px !important;
  }
  .top-version-badge {font-size:12px !important; padding:4px 9px !important;}
  .main-title {font-size:24px !important; line-height:1.2 !important;}
  .sub-title {font-size:12px !important; line-height:1.35 !important;}
}

</style>
""", unsafe_allow_html=True)

st.markdown(
    f'''
    <div class="top-version-card">
        <div class="top-version-badge">최신 버전 · V37</div>
        <div class="main-title">📖 {APP_TITLE}</div>
        <div class="sub-title">{APP_VERSION} | AI 공동작가 · 작품 생성 · 수정 · 보관함 이어쓰기 · 자동저장 · 작품 보관함 중심</div>
    </div>
    ''',
    unsafe_allow_html=True
)


def now_str():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def slugify(name: str) -> str:
    name = (name or "inyoung_project").strip()
    safe = re.sub(r"[^가-힣A-Za-z0-9_.-]+", "_", name)
    return safe[:80] or "inyoung_project"


def new_project_template(name: str = "신의 계시") -> dict:
    return {
        "project_name": name,
        "created_at": now_str(),
        "updated_at": now_str(),
        "current_version": 0,
        "status": "작업중",
        "master_text": "",
        "world_bible": "",
        "characters": "",
        "power_system": "",
        "episode_outline": "",
        "latest_draft": "",
        "latest_result": "",
        "last_action": "",
        "uploaded_name": "",
        "uploaded_text": "",
        "versions": [],
        "timeline": [],
    }


DEFAULTS = {
    "current_project_name_v30": "신의 계시",
    "project_v30": new_project_template("신의 계시"),
    "openai_input_tokens_v30": 0,
    "openai_output_tokens_v30": 0,
    "openai_total_tokens_v30": 0,
    "openai_estimated_cost_v30": 0.0,
    "openai_call_count_v30": 0,
    "openai_actual_balance_v30": None,
    "last_error_v30": "",
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = json.loads(json.dumps(v, ensure_ascii=False)) if isinstance(v, dict) else v


def project_path(project_name: str | None = None) -> Path:
    p = st.session_state.project_v30
    name = project_name or p.get("project_name", "신의 계시")
    return PROJECT_DIR / f"{slugify(name)}.json"


def save_index():
    projects = []
    for file in PROJECT_DIR.glob("*.json"):
        if file.name.startswith("_"):
            continue
        try:
            data = json.loads(file.read_text(encoding="utf-8"))
            projects.append({
                "project_name": data.get("project_name", file.stem),
                "updated_at": data.get("updated_at", ""),
                "current_version": data.get("current_version", 0),
                "status": data.get("status", "작업중"),
                "file": file.name,
            })
        except Exception:
            pass
    projects.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
    INDEX_FILE.write_text(json.dumps(projects, ensure_ascii=False, indent=2), encoding="utf-8")


def load_index():
    save_index()
    if not INDEX_FILE.exists():
        return []
    try:
        return json.loads(INDEX_FILE.read_text(encoding="utf-8"))
    except Exception:
        return []


def save_project():
    p = st.session_state.project_v30
    p["updated_at"] = now_str()
    project_path().write_text(json.dumps(p, ensure_ascii=False, indent=2), encoding="utf-8")
    save_index()


def load_project(name: str) -> bool:
    path = project_path(name)
    if not path.exists():
        return False
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        st.session_state.project_v30 = data
        st.session_state.current_project_name_v30 = data.get("project_name", name)
        return True
    except Exception as e:
        st.session_state.last_error_v30 = f"프로젝트 불러오기 실패: {e}"
        return False


def add_timeline(action: str, detail: str):
    p = st.session_state.project_v30
    p.setdefault("timeline", [])
    p["timeline"].insert(0, {
        "time": now_str(),
        "version": p.get("current_version", 0),
        "action": action,
        "detail": (detail or "")[:1000],
    })
    p["timeline"] = p["timeline"][:200]


def export_current_text() -> str:
    p = st.session_state.project_v30
    return f"""# {p.get('project_name','인영이 작품')} - AI 공동작가 프로젝트 V37

[현재 버전]
V{p.get('current_version', 0)}

[상태]
{p.get('status','작업중')}

[최근 작업]
{p.get('last_action','')}

[마스터 원고]
{p.get('master_text','')}

[세계관]
{p.get('world_bible','')}

[등장인물]
{p.get('characters','')}

[세력/무공/능력]
{p.get('power_system','')}

[회차 구성]
{p.get('episode_outline','')}

[최근 원고]
{p.get('latest_draft','')}

[최근 AI 결과]
{p.get('latest_result','')}

[타임라인]
{json.dumps(p.get('timeline', []), ensure_ascii=False, indent=2)}
"""


def auto_save_version(action: str, detail: str, content: str | None = None):
    p = st.session_state.project_v30
    p.setdefault("versions", [])
    version_no = int(p.get("current_version", 0)) + 1
    p["current_version"] = version_no
    p["updated_at"] = now_str()
    p["last_action"] = action
    text = content if content is not None else export_current_text()
    snapshot = {
        "version_no": version_no,
        "label": f"V{version_no} {action}",
        "action": action,
        "detail": (detail or "")[:1200],
        "saved_at": now_str(),
        "content": text or "",
        "master_text": p.get("master_text", ""),
        "world_bible": p.get("world_bible", ""),
        "characters": p.get("characters", ""),
        "power_system": p.get("power_system", ""),
        "episode_outline": p.get("episode_outline", ""),
        "latest_draft": p.get("latest_draft", ""),
        "latest_result": p.get("latest_result", ""),
    }
    p["versions"].insert(0, snapshot)
    p["versions"] = p["versions"][:300]
    add_timeline(action, detail)
    save_project()
    return snapshot


def create_new_project(name: str):
    st.session_state.project_v30 = new_project_template(name)
    st.session_state.current_project_name_v30 = name
    auto_save_version("새 작품 생성", "새 작품을 생성했습니다.", "")
    save_project()


if not st.session_state.project_v30.get("versions"):
    if project_path(st.session_state.current_project_name_v30).exists():
        load_project(st.session_state.current_project_name_v30)


def current_balance() -> float:
    if st.session_state.openai_actual_balance_v30 is not None:
        return max(0.0, float(st.session_state.openai_actual_balance_v30))
    return max(0.0, OPENAI_INITIAL_CREDIT - float(st.session_state.openai_estimated_cost_v30))


def add_usage(response):
    usage = getattr(response, "usage", None)
    in_tok = int(getattr(usage, "input_tokens", 0) or getattr(usage, "prompt_tokens", 0) or 0) if usage else 0
    out_tok = int(getattr(usage, "output_tokens", 0) or getattr(usage, "completion_tokens", 0) or 0) if usage else 0
    total = int(getattr(usage, "total_tokens", 0) or (in_tok + out_tok)) if usage else 0
    cost = (in_tok / 1_000_000 * OPENAI_INPUT_COST_PER_1M) + (out_tok / 1_000_000 * OPENAI_OUTPUT_COST_PER_1M)
    st.session_state.openai_input_tokens_v30 += in_tok
    st.session_state.openai_output_tokens_v30 += out_tok
    st.session_state.openai_total_tokens_v30 += total
    st.session_state.openai_estimated_cost_v30 += cost
    st.session_state.openai_call_count_v30 += 1
    if st.session_state.openai_actual_balance_v30 is not None:
        st.session_state.openai_actual_balance_v30 = max(0.0, float(st.session_state.openai_actual_balance_v30) - cost)


def model_candidates():
    base = [OPENAI_MODEL, "gpt-4.1-mini", "gpt-4o-mini", "gpt-4.1"]
    result = []
    for m in base:
        if m and m not in result:
            result.append(m)
    return result


def openai_text(prompt: str, fallback: str = "", temperature: float = 0.55, max_output_tokens: int = 7000) -> str:
    if not OPENAI_API_KEY or OpenAI is None:
        return fallback or "OPENAI_API_KEY가 없어 기본 초안만 표시합니다."
    client = OpenAI(api_key=OPENAI_API_KEY)
    last_error = None
    for model in model_candidates():
        try:
            response = client.responses.create(model=model, input=prompt, temperature=temperature, max_output_tokens=max_output_tokens)
            add_usage(response)
            text = getattr(response, "output_text", "")
            if text:
                st.session_state.last_error_v30 = ""
                return text
        except Exception as e:
            last_error = e
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "너는 한국 웹소설, 드라마, 장르문학을 함께 개발하는 전문 AI 공동작가다. 사용자의 핵심 설정을 보존하고 오리지널 세계관으로 발전시킨다."},
                    {"role": "user", "content": prompt},
                ],
                temperature=temperature,
                max_tokens=max_output_tokens,
            )
            add_usage(response)
            text = response.choices[0].message.content
            if text:
                st.session_state.last_error_v30 = ""
                return text
        except Exception as e:
            last_error = e
    st.session_state.last_error_v30 = str(last_error)
    return f"OpenAI 생성 실패: {last_error}\n\n기본 초안:\n{fallback}"


def estimate_cost(chars: int, mode: str) -> float:
    base = max(chars, 500)
    if mode == "small":
        return min(0.12, base / 50000 * 0.06 + 0.01)
    if mode == "medium":
        return min(0.32, base / 40000 * 0.13 + 0.03)
    if mode == "large":
        return min(0.90, base / 30000 * 0.26 + 0.10)
    return min(0.25, base / 40000 * 0.10 + 0.02)


def read_uploaded_file(uploaded):
    if uploaded is None:
        return ""
    name = uploaded.name.lower()
    if name.endswith(".txt"):
        raw = uploaded.read()
        for enc in ["utf-8", "cp949", "euc-kr"]:
            try:
                return raw.decode(enc)
            except Exception:
                pass
        return raw.decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        if PdfReader is None:
            return "PDF 분석을 위해 pypdf가 필요합니다."
        reader = PdfReader(uploaded)
        pages = []
        for page in reader.pages[:120]:
            try:
                pages.append(page.extract_text(extraction_mode="layout") or "")
            except Exception:
                pages.append(page.extract_text() or "")
        return "\n".join(pages)
    if name.endswith(".docx"):
        if Document is None:
            return "DOCX 분석을 위해 python-docx가 필요합니다."
        doc = Document(uploaded)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return "지원 형식: TXT, PDF, DOCX"


def make_docx_bytes(title: str, body: str):
    if Document is None:
        return None
    bio = io.BytesIO()
    doc = Document()
    doc.add_heading(title, level=1)
    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("#"):
            doc.add_heading(block.strip("# ").strip(), level=2)
        elif block.startswith("[") and block.endswith("]"):
            doc.add_heading(block.strip("[]"), level=2)
        else:
            doc.add_paragraph(block)
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def make_project_zip_bytes(project: dict):
    """작품 전체를 ZIP으로 백업합니다."""
    import zipfile
    bio = io.BytesIO()
    base = slugify(project.get("project_name", "inyoung_project"))
    with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{base}_project.json", json.dumps(project, ensure_ascii=False, indent=2))
        zf.writestr(f"{base}_current.txt", export_current_text())
        for v in project.get("versions", []):
            fname = slugify(f"V{v.get('version_no')}_{v.get('action','version')}")
            zf.writestr(f"versions/{fname}.txt", v.get("content", ""))
            zf.writestr(f"versions/{fname}.json", json.dumps(v, ensure_ascii=False, indent=2))
    bio.seek(0)
    return bio.getvalue()


def delete_version_by_no(version_no: int) -> bool:
    """선택 버전을 삭제하고 프로젝트를 저장합니다."""
    p = st.session_state.project_v30
    before = len(p.get("versions", []))
    p["versions"] = [v for v in p.get("versions", []) if int(v.get("version_no", -1)) != int(version_no)]
    after = len(p.get("versions", []))
    if after < before:
        add_timeline("버전 삭제", f"V{version_no} 버전을 삭제했습니다.")
        save_project()
        return True
    return False


def update_version_by_no(version_no: int, new_content: str) -> bool:
    """선택 버전 자체를 직접 수정 저장합니다. V1/V2는 기본 보호 옵션을 UI에서 확인합니다."""
    p = st.session_state.project_v30
    for v in p.get("versions", []):
        if int(v.get("version_no", -1)) == int(version_no):
            v["content"] = new_content
            v["master_text"] = new_content
            v["latest_result"] = new_content
            v["latest_draft"] = new_content
            v["edited_at"] = now_str()
            v["detail"] = (v.get("detail", "") + "\n[직접편집] 작품보관함에서 선택 버전 내용을 직접 수정했습니다.").strip()[:1200]
            add_timeline("선택 버전 직접 편집", f"V{version_no} 내용을 직접 수정 저장했습니다.")
            save_project()
            return True
    return False


def project_context(max_chars: int = 32000) -> str:
    p = st.session_state.project_v30
    parts = [
        f"[작품명]\n{p.get('project_name','')}",
        f"[마스터 원고]\n{p.get('master_text','')}",
        f"[세계관]\n{p.get('world_bible','')}",
        f"[등장인물]\n{p.get('characters','')}",
        f"[세력/무공/능력]\n{p.get('power_system','')}",
        f"[회차 구성]\n{p.get('episode_outline','')}",
        f"[최근 원고]\n{p.get('latest_draft','')}",
        f"[최근 AI 결과]\n{p.get('latest_result','')}",
    ]
    return "\n\n".join(parts)[:max_chars]


def original_rule(user_text: str = "") -> str:
    return f"""
저작권/오리지널 규칙:
- 특정 작품명, 작가명, 캐릭터명, 설정을 예시로 언급해도 이름/설정/줄거리/고유 용어를 복사하지 않는다.
- 예시는 독자가 원하는 재미, 분위기, 장르적 방향으로만 해석한다.
- 최종 결과는 새로운 제목, 인물명, 세력명, 무공명, 세계관, 사건으로 만든다.
- 사용자가 직접 만든 핵심 설정은 최대한 보존하되, 완성도를 높이는 방향으로 확장한다.
사용자 요청: {user_text}
"""


def ai_create_project(idea: str, genre: str, tone: str, length_level: str) -> str:
    prompt = f"""
너는 AI 공동작가 플랫폼의 메인 작가다.
사용자의 한 문장 아이디어를 장기 연재 가능한 작품으로 확장해라.

{original_rule(idea)}

목표 장르: {genre}
톤/재미 방향: {tone}
분량 방향: {length_level}

반드시 아래 형식으로 작성:
# 작품명 후보
# 한 줄 로그라인
# 작품 콘셉트
# 세계관
# 등장인물
# 세력/무공/능력 체계
# 전체 줄거리
# 회차 확장 방향
# 독자 후킹 장치
# 앞으로 수정/확장할 때 지켜야 할 핵심 설정
"""
    return openai_text(prompt, "작품 생성 실패", temperature=0.62, max_output_tokens=9000)


def ai_modify_work(user_request: str, source_text: str = "") -> str:
    prompt = f"""
너는 AI 공동작가이자 원고 수정 전문 편집자다.
아래 현재 작품 정보와 사용자의 수정 요청을 바탕으로 수정본을 작성해라.

{original_rule(user_request)}

현재 작품 정보:
{project_context(26000)}

업로드/선택 원고:
{source_text[:30000]}

사용자 수정 요청:
{user_request}

해야 할 일:
1. 수정 요청 핵심 요약
2. 기존 설정과 충돌 여부 확인
3. 수정 방향 제안
4. 수정본 작성
5. 이후 이어질 내용 제안
"""
    return openai_text(prompt, "수정본 생성 실패", temperature=0.55, max_output_tokens=12000)


def ai_continue_episode(request: str, episode_no: int, style: str) -> str:
    prompt = f"""
너는 장기 연재 웹소설의 메인 작가다.
현재 작품 DB를 기반으로 {episode_no}화를 이어서 작성해라.

현재 작품 정보:
{project_context(32000)}

작성 요청:
{request}

문체/연출 방향:
{style}

결과 형식:
# {episode_no}화 제목
본문
[다음 화 후킹]
"""
    return openai_text(prompt, "이어쓰기 실패", temperature=0.72, max_output_tokens=13000)


def ai_continue_from_selected_versions(selected_text: str, request: str, episode_no: int, style: str) -> str:
    """V37: 작품보관함에서 선택한 버전을 기준으로 이어쓰기를 진행합니다."""
    prompt = f"""
너는 장기 연재 웹소설의 메인 작가다.
사용자가 작품보관함에서 선택한 버전을 기준으로 다음 회차를 이어서 작성해라.

중요 규칙:
- 선택한 버전의 설정, 사건, 문체를 최우선으로 유지한다.
- 기존 내용과 충돌하지 않도록 이어쓴다.
- 급격한 파워업보다 독자가 납득할 수 있는 성장 과정과 갈등을 넣는다.
- 결과는 새 버전으로 저장될 수 있게 완성본 형태로 작성한다.

선택한 작품보관함 버전 내용:
{selected_text[:36000]}

작성할 회차: {episode_no}화
문체/연출 방향: {style}
사용자 요청:
{request}

결과 형식:
# {episode_no}화 제목
본문
[이번 화 핵심 변화]
[다음 화 후킹]
"""
    return openai_text(prompt, "선택 버전 기준 이어쓰기 실패", temperature=0.72, max_output_tokens=13000)


def ai_story_analysis_continue(source_text: str, genre_choices: list, output_type: str, continue_range: str, user_request: str, tone_level: str) -> str:
    """V37: 기존 구성안/시놉/원고를 분석하고, 선택한 장르 방향으로 이후 구성안을 이어서 작성합니다."""
    genres = ", ".join(genre_choices) if genre_choices else "사용자 지정 장르 없음"
    prompt = f"""
너는 드라마/웹소설/영화 구성안을 분석하고 이어 쓰는 전문 AI 스토리 개발자다.
사용자가 이전에 작성한 구성안, 시놉시스, 원고를 먼저 분석한 뒤,
선택한 장르 요소를 반영하여 이후 내용을 새롭게 확장해라.

핵심 규칙:
- 기존 구성안의 핵심 설정, 인물 관계, 사건 흐름을 먼저 분석한다.
- 기존 내용을 무시하고 새로 만들지 않는다.
- 사용자가 선택한 장르 요소를 자연스럽게 섞는다.
- 스릴러/반전/코믹/감동/액션/로맨스/미스터리 등은 억지로 붙이지 말고 사건 구조 안에 녹인다.
- 이후 작성할 내용은 실제 작가가 바로 이어 쓸 수 있게 구체적인 구성안 형태로 작성한다.
- 기존 작품명/인물명이 있으면 유지하되, 충돌되거나 어색한 이름은 대안도 제시한다.
- 결과는 새 버전으로 저장될 수 있게 완성형으로 작성한다.

선택 장르:
{genres}

결과 형식:
{output_type}

이어쓸 범위:
{continue_range}

강도/톤:
{tone_level}

사용자 추가 요청:
{user_request}

분석할 기존 구성안/시놉/원고:
{source_text[:38000]}

반드시 아래 형식으로 작성:
# AI 스토리 분석 결과
## 1. 기존 구성안 핵심 요약
## 2. 현재까지의 인물/관계/갈등 분석
## 3. 선택 장르 반영 방향
## 4. 이후 전개 핵심 아이디어
## 5. 이어질 구성안
## 6. 반전/코믹/스릴러/감동 포인트
## 7. 다음 작업 추천
"""
    return openai_text(prompt, "스토리 분석 및 이어쓰기 실패", temperature=0.68, max_output_tokens=14000)


def ai_cowriter_command(command: str, mode: str) -> str:
    """V31: 마스터 원고를 직접 덮어쓰지 않고, 사용자의 지시를 작품 DB 위에 반영한 새 결과를 생성합니다."""
    prompt = f"""
너는 이 작품을 장기 연재로 함께 쓰는 AI 공동작가다.
아래 현재 작품 DB와 마스터 원고를 절대 잃어버리지 말고, 사용자의 지시를 반영해 새로운 작업 결과를 작성해라.

중요 저장 규칙:
- 기존 V1, V2, 이전 버전 내용은 삭제하지 않는다.
- 기존 마스터 원고를 함부로 축약하거나 덮어쓰지 않는다.
- 사용자의 지시는 새 작업 결과로 생성한다.
- 이전 설정과 충돌하면 '충돌 지점'과 '해결안'을 먼저 설명한다.
- 결과는 작품 보관함에 새 버전으로 저장될 수 있게 완성된 형태로 작성한다.

작업 모드:
{mode}

현재 작품 DB:
{project_context(36000)}

사용자 지시:
{command}

반드시 아래 형식으로 작성:
# AI 공동작가 작업 결과
## 1. 지시 해석
## 2. 기존 설정 보존 확인
## 3. 충돌/보완 지점
## 4. 최종 결과물
## 5. 다음 작업 추천
"""
    return openai_text(prompt, "AI 공동작가 지시 실행 실패", temperature=0.62, max_output_tokens=14000)


with st.sidebar:
    st.header("📚 작품 선택")
    projects = load_index()
    names = [x["project_name"] for x in projects]
    current_name = st.session_state.project_v30.get("project_name", "신의 계시")
    if names:
        idx = names.index(current_name) if current_name in names else 0
        selected_name = st.selectbox("저장된 작품", names, index=idx)
        if selected_name != current_name:
            if load_project(selected_name):
                st.rerun()
    else:
        st.caption("저장된 작품이 없습니다.")

    new_name = st.text_input("새 작품명", value="")
    if st.button("➕ 새 작품 만들기"):
        if new_name.strip():
            create_new_project(new_name.strip())
            st.success("새 작품 생성 완료")
            st.rerun()
        else:
            st.warning("작품명을 입력해 주세요.")

    st.divider()
    p = st.session_state.project_v30
    p["project_name"] = st.text_input("현재 작품명 수정", value=p.get("project_name", "신의 계시"))
    if st.button("💾 작품명/현재상태 저장"):
        save_project()
        st.success("저장 완료")
        st.rerun()

    st.divider()
    st.caption("V37는 AI 작업 성공 시 자동저장됩니다.")

p = st.session_state.project_v30
bal = current_balance()
spent = OPENAI_INITIAL_CREDIT - bal if st.session_state.openai_actual_balance_v30 is not None else st.session_state.openai_estimated_cost_v30

m1, m2, m3 = st.columns(3)
m1.metric("📖 현재 작품", p.get("project_name", "-"))
m2.metric("현재 버전", f"V{p.get('current_version', 0)}")
m3.metric("최근 저장", p.get("updated_at", "-")[-8:] if p.get("updated_at") else "-")

if st.session_state.last_error_v30:
    st.warning(f"최근 OpenAI 오류: {st.session_state.last_error_v30}")

st.markdown(f"""
<div class='notice'>
<b>자동저장 상태:</b> AI 작업이 성공하면 자동으로 작품 보관함에 버전이 생성됩니다.
현재 작품 <b>{p.get('project_name','-')}</b> / 현재 버전 <b>V{p.get('current_version', 0)}</b>
</div>
""", unsafe_allow_html=True)

tabs = st.tabs(["🏠 내 작품", "✨ 작품 생성", "✏️ 작품 수정", "📚 작품 보관함", "📥 다운로드", "📊 OpenAI", "🎬 스토리 분석실", "🤖 AI 공동작가", "⚙️ 가이드"])

with tabs[0]:
    st.subheader("🏠 내 작품")
    p = st.session_state.project_v30
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("작품명", p.get("project_name", "-"))
    c2.metric("버전", f"V{p.get('current_version', 0)}")
    c3.metric("보관 버전", f"{len(p.get('versions', []))}개")
    c4.metric("상태", p.get("status", "작업중"))

    st.markdown("### 최근 작업")
    if p.get("latest_result"):
        st.markdown(f"<div class='ai-box'>{p.get('latest_result')[:5000]}</div>", unsafe_allow_html=True)
    elif p.get("master_text"):
        st.markdown(f"<div class='ai-box'>{p.get('master_text')[:5000]}</div>", unsafe_allow_html=True)
    else:
        st.info("아직 작업 결과가 없습니다. '작품 생성' 탭에서 시작하세요.")

    st.markdown("### 작업 타임라인")
    if p.get("timeline"):
        st.dataframe(pd.DataFrame(p.get("timeline", [])), use_container_width=True, hide_index=True)
    else:
        st.info("아직 타임라인이 없습니다.")

    st.markdown("### 마스터 원고 직접 편집")
    st.caption("직접 수정한 뒤 저장하면 새 버전으로 보관됩니다. 저장 전에도 TXT/DOCX로 바로 다운로드할 수 있습니다.")
    p["master_text"] = st.text_area("마스터 원고", value=p.get("master_text", ""), height=420)
    master_name = slugify(f"{p.get('project_name','작품')}_master_V{p.get('current_version',0)}")
    mc1, mc2, mc3, mc4 = st.columns(4)
    if mc1.button("💾 마스터 원고 직접 저장"):
        auto_save_version("마스터 원고 직접 저장", "사용자가 마스터 원고를 직접 수정하여 저장했습니다.", p.get("master_text", ""))
        st.success("자동 버전 저장 완료")
        st.rerun()
    mc2.download_button("📥 마스터 TXT", p.get("master_text", ""), f"{master_name}.txt", "text/plain")
    master_docx = make_docx_bytes(f"{p.get('project_name','작품')} 마스터 원고", p.get("master_text", ""))
    if master_docx:
        mc3.download_button("📥 마스터 DOCX", master_docx, f"{master_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        mc3.info("DOCX 사용 불가")
    if mc4.button("🗂 보관함 새 버전 저장"):
        snap = auto_save_version("마스터 원고 보관함 저장", "현재 마스터 원고를 작품보관함에 새 버전으로 저장했습니다.", p.get("master_text", ""))
        st.success(f"작품보관함 저장 완료: {snap.get('label')}")
        st.rerun()

with tabs[1]:
    st.subheader("✨ 작품 생성")
    idea = st.text_area("내가 만들고 싶은 작품을 편하게 적어주세요", height=160, placeholder="예: 판타지 무협소설을 쓰고 싶어. 천마신교 5번째 제자가 현대 헌터 세계의 전생을 깨닫고, 주신의 의뢰로 천마가 된 뒤 판타지 세계를 구하는 먼치킨 장편으로 만들고 싶어.")
    c1, c2, c3 = st.columns(3)
    genre = c1.selectbox("목표 장르", ["판타지 무협", "현대판타지", "헌터물", "먼치킨", "무협", "다크판타지", "로맨스판타지", "SF"], index=0)
    tone = c2.selectbox("톤", ["초강력 먼치킨", "왕도 성장물", "비장한 영웅담", "다크 판타지", "코믹 액션", "신화적 대서사"], index=0)
    length_level = c3.selectbox("분량", ["단편", "중편", "장편", "초장편 연재"], index=2)
    if st.button("🚀 작품 생성 + 자동저장", type="primary"):
        if not idea.strip():
            st.error("아이디어를 먼저 입력해 주세요.")
        else:
            prog = st.progress(0); status = st.empty()
            status.write("20% · 아이디어 분석 중..."); prog.progress(20)
            result = ai_create_project(idea, genre, tone, length_level)
            status.write("75% · 작품 DB 정리 및 자동저장 중..."); prog.progress(75)
            p = st.session_state.project_v30
            p["master_text"] = result; p["latest_result"] = result; p["world_bible"] = result; p["last_action"] = "작품 생성"
            snap = auto_save_version("작품 생성", idea, result)
            status.write(f"100% · 자동저장 완료: {snap.get('label')}"); prog.progress(100)
            st.success(f"작품 생성 완료 및 자동저장 완료: {snap.get('label')}")
            st.rerun()

with tabs[2]:
    st.subheader("✏️ 작품 수정")
    source_mode = st.radio("수정 대상", ["현재 저장된 작품", "파일 업로드 원고"], horizontal=True)
    source_text = ""
    if source_mode == "파일 업로드 원고":
        uploaded = st.file_uploader("TXT/PDF/DOCX 업로드", type=["txt", "pdf", "docx"], key="upload_v30_modify")
        if uploaded is not None:
            if st.button("📖 업로드 파일 읽기"):
                text = read_uploaded_file(uploaded)
                p = st.session_state.project_v30
                p["uploaded_name"] = uploaded.name; p["uploaded_text"] = text
                save_project()
                st.success(f"업로드 저장 완료: {uploaded.name} / 글자수 {len(text):,}")
                st.rerun()
        if p.get("uploaded_text"):
            st.success(f"현재 업로드 원고: {p.get('uploaded_name','')} / 글자수 {len(p.get('uploaded_text','')):,}")
            source_text = p.get("uploaded_text", "")
    else:
        source_text = export_current_text()
    request = st.text_area("어떻게 수정하고 싶은지 적어주세요", height=150, placeholder="예: 주인공을 더 압도적인 먼치킨으로 만들고, 천마신교 5번째 제자 설정을 더 강하게 살려줘. 전투씬은 더 강하게, 감동과 반전도 추가해줘.")
    if st.button("🤖 AI 수정 + 자동저장", type="primary"):
        if not request.strip():
            st.error("수정 요청을 입력해 주세요.")
        elif not source_text.strip():
            st.error("수정할 원고가 없습니다.")
        else:
            prog = st.progress(0); status = st.empty()
            status.write("25% · 현재 작품/업로드 원고 분석 중..."); prog.progress(25)
            result = ai_modify_work(request, source_text)
            status.write("80% · 수정본 자동저장 중..."); prog.progress(80)
            p = st.session_state.project_v30
            p["latest_result"] = result; p["latest_draft"] = result; p["master_text"] = result
            snap = auto_save_version("작품 수정", request, result)
            status.write(f"100% · 자동저장 완료: {snap.get('label')}"); prog.progress(100)
            st.success(f"수정 완료 및 자동저장 완료: {snap.get('label')}")
            st.rerun()
    if p.get("latest_result"):
        st.markdown("### 최근 수정/생성 결과")
        st.markdown(f"<div class='ai-box'>{p.get('latest_result')[:8000]}</div>", unsafe_allow_html=True)

with tabs[3]:
    st.subheader("📚 작품 보관함")
    st.markdown("""
<div class='card'>
<b>V37 작품보관함</b><br>
버전 왼쪽의 선택란을 체크한 뒤 <b>다운로드 / 삭제 / AI 수정 / 이어쓰기</b>를 바로 실행할 수 있습니다.<br>
기존 별도 이어쓰기 탭은 제거하고, <b>선택한 버전 기준으로 이어쓰기</b>를 진행하도록 통합했습니다. <b>V1/V2는 기본 보호</b>됩니다.
</div>
""", unsafe_allow_html=True)

    versions = p.get("versions", [])
    if not versions:
        st.info("아직 저장된 버전이 없습니다.")
    else:
        rows = []
        for vrow in versions:
            rows.append({
                "선택": False,
                "번호": int(vrow.get("version_no", 0)),
                "버전": f"V{vrow.get('version_no')}",
                "작업": vrow.get("action", ""),
                "저장시간": vrow.get("saved_at", ""),
                "글자수": len(vrow.get("content", "")),
                "설명": vrow.get("detail", "")[:70],
            })

        st.markdown("### ✅ 작업할 버전 선택")
        edited_df = st.data_editor(
            pd.DataFrame(rows),
            use_container_width=True,
            hide_index=True,
            disabled=["번호", "버전", "작업", "저장시간", "글자수", "설명"],
            column_config={
                "선택": st.column_config.CheckboxColumn("선택", help="다운로드/삭제/수정할 버전을 선택하세요."),
                "번호": st.column_config.NumberColumn("번호", width="small"),
                "버전": st.column_config.TextColumn("버전", width="small"),
                "작업": st.column_config.TextColumn("작업", width="medium"),
                "저장시간": st.column_config.TextColumn("저장시간", width="medium"),
                "글자수": st.column_config.NumberColumn("글자수", width="small"),
                "설명": st.column_config.TextColumn("설명", width="large"),
            },
            key="version_select_editor_v35",
        )

        selected_nos = []
        try:
            selected_nos = [int(r["번호"]) for _, r in edited_df.iterrows() if bool(r.get("선택"))]
        except Exception:
            selected_nos = []

        if not selected_nos:
            selected_nos = [int(versions[0].get("version_no", 0))]
            st.caption("선택한 버전이 없어 최신 버전을 기본으로 보여줍니다.")

        selected_versions = [vv for vv in versions if int(vv.get("version_no", -1)) in selected_nos]
        primary_v = selected_versions[0]
        version_no = int(primary_v.get("version_no", 0))
        safe_name = slugify(f"{p.get('project_name','작품')}_{primary_v.get('label','version')}")

        st.markdown(f"### 선택 상태: {', '.join(['V'+str(x) for x in selected_nos])}")

        st.markdown("### 📥 선택 버전 다운로드")
        dc1, dc2, dc3 = st.columns(3)
        if len(selected_versions) == 1:
            dc1.download_button("📥 선택 TXT", primary_v.get("content", ""), f"{safe_name}.txt", "text/plain")
            dc2.download_button("📥 선택 JSON", json.dumps(primary_v, ensure_ascii=False, indent=2), f"{safe_name}.json", "application/json")
            docx_data = make_docx_bytes(primary_v.get("label", "version"), primary_v.get("content", ""))
            if docx_data:
                dc3.download_button("📥 선택 DOCX", docx_data, f"{safe_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                dc3.info("DOCX 사용 불가")
        else:
            import zipfile
            zip_bio = io.BytesIO()
            with zipfile.ZipFile(zip_bio, "w", zipfile.ZIP_DEFLATED) as zf:
                for sv in selected_versions:
                    fname = slugify(f"V{sv.get('version_no')}_{sv.get('action','version')}")
                    zf.writestr(f"{fname}.txt", sv.get("content", ""))
                    zf.writestr(f"{fname}.json", json.dumps(sv, ensure_ascii=False, indent=2))
            zip_bio.seek(0)
            dc1.download_button("📦 선택 버전 ZIP", zip_bio.getvalue(), f"{slugify(p.get('project_name','작품'))}_selected_versions.zip", "application/zip")
            dc2.info("여러 버전 선택 시 ZIP으로 다운로드됩니다.")
            dc3.info("DOCX는 단일 버전 선택 시 제공됩니다.")

        st.markdown("### 👁 선택 버전 미리보기")
        preview_limit = st.slider("미리보기 글자수", 800, 8000, 2500, 200, key="preview_limit_v35")
        st.markdown(f"<div class='memo-box compact-preview'>{primary_v.get('content','')[:preview_limit]}</div>", unsafe_allow_html=True)

        st.markdown("### ✏️ 선택 버전 직접 편집")
        if len(selected_versions) != 1:
            st.info("직접 편집은 한 번에 한 개 버전만 가능합니다. 하나만 체크해 주세요.")
        else:
            st.caption("기존 버전을 유지하려면 '편집본을 새 버전으로 저장'을 사용하세요. V1/V2는 원본 보호됩니다.")
            edit_key = f"version_edit_v35_{version_no}"
            edited_content = st.text_area("선택 버전 편집 내용", value=primary_v.get("content", ""), height=340, key=edit_key)
            ec1, ec2, ec3 = st.columns(3)
            if ec1.button("💾 편집본 새 버전 저장", key=f"save_new_from_edit_v35_{version_no}"):
                p["latest_result"] = edited_content
                p["latest_draft"] = edited_content
                p["master_text"] = edited_content
                snap = auto_save_version(f"V{version_no} 편집본 저장", f"{primary_v.get('label')} 편집본을 새 버전으로 저장했습니다.", edited_content)
                st.success(f"편집본 새 버전 저장 완료: {snap.get('label')}")
                st.rerun()

            direct_protect = version_no <= 2
            direct_unlock = True
            if direct_protect:
                direct_unlock = st.checkbox("V1/V2 직접 편집 보호 해제", key=f"unlock_direct_edit_v35_{version_no}")

            if ec2.button("✏️ 선택 버전 자체 수정 저장", key=f"direct_save_selected_version_v35_{version_no}"):
                if direct_protect and not direct_unlock:
                    st.error("V1/V2 직접 편집 보호 해제가 필요합니다. 안전하게는 새 버전 저장을 권장합니다.")
                else:
                    if update_version_by_no(version_no, edited_content):
                        st.success(f"V{version_no} 선택 버전 자체를 수정 저장했습니다.")
                        st.rerun()
                    else:
                        st.error("선택 버전을 찾지 못했습니다.")

            if ec3.button("🗂 편집본을 마스터에 반영", key=f"apply_master_only_v35_{version_no}"):
                p["master_text"] = edited_content
                p["latest_result"] = edited_content
                p["latest_draft"] = edited_content
                save_project()
                st.success("마스터 원고에 반영했습니다. 보관함에 남기려면 새 버전 저장을 눌러주세요.")
                st.rerun()

        st.markdown("### 🤖 선택 버전 AI 수정 요청")
        st.caption("체크한 버전 내용을 AI가 읽고, 요청한 방향으로 수정본을 새 버전으로 저장합니다.")
        ai_req = st.text_area(
            "AI에게 수정할 부분 요청",
            height=110,
            placeholder="예: 선택한 버전에서 주인공의 전투 장면을 더 압도적으로 만들고, 감동과 반전을 추가해줘.",
            key="selected_version_ai_request_v35"
        )
        ai_mode = st.selectbox("AI 수정 방식", ["선택 버전 전체 수정", "설정/세계관 정리", "문장 다듬기", "전투씬 강화", "감정선 강화", "장기 플롯 확장"], key="selected_version_ai_mode_v35")
        selected_text = "\n\n\n--- 선택 버전 구분 ---\n\n\n".join([
            f"# V{sv.get('version_no')} · {sv.get('action','')}\n\n{sv.get('content','')}"
            for sv in selected_versions
        ])
        if st.button("🤖 선택 버전 AI 수정 + 새 버전 저장", type="primary", key="ai_modify_selected_versions_v35"):
            if not ai_req.strip():
                st.error("AI에게 요청할 수정 내용을 입력해 주세요.")
            else:
                prog = st.progress(0); status = st.empty()
                status.write("25% · 선택 버전 내용 분석 중..."); prog.progress(25)
                request_full = f"[작업방식: {ai_mode}]\n{ai_req}"
                result = ai_modify_work(request_full, selected_text)
                status.write("80% · AI 수정본 새 버전 저장 중..."); prog.progress(80)
                p["latest_result"] = result
                p["latest_draft"] = result
                p["master_text"] = result
                snap = auto_save_version(f"선택 버전 AI 수정 · {ai_mode}", f"선택 버전 {selected_nos} 수정 요청: {ai_req}", result)
                status.write(f"100% · 저장 완료: {snap.get('label')}"); prog.progress(100)
                st.success(f"선택 버전 AI 수정 완료 및 새 버전 저장: {snap.get('label')}")
                st.rerun()

        st.markdown("### 📖 선택 버전 기준 이어쓰기")
        st.caption("체크한 버전을 기준으로 다음 회차를 이어씁니다. 기존 별도 이어쓰기 탭보다 현재 작업 버전 기준이 명확합니다.")
        cw1, cw2 = st.columns(2)
        continue_episode_no = cw1.number_input("작성할 회차", min_value=1, max_value=2000, value=1, step=1, key="continue_episode_no_v36")
        continue_style = cw2.selectbox("문체/연출", ["웹소설 빠른 몰입형", "무협지 고풍 문체", "현대판타지 경쾌한 문체", "다크하고 비장한 문체", "전투씬 강화", "감정선 강화", "반전 강화", "성장 서사 강화"], index=0, key="continue_style_v36")
        continue_req = st.text_area(
            "이어쓰기 요청",
            height=110,
            placeholder="예: 선택한 버전 다음 내용으로 10화를 써줘. 초반 전개가 너무 빠르지 않게 하고, 주인공의 고생과 성장 과정이 보이게 해줘.",
            key="continue_request_selected_v36"
        )
        if st.button("📖 선택 버전 기준 이어쓰기 + 새 버전 저장", type="primary", key="continue_selected_versions_v36"):
            if not continue_req.strip():
                st.error("이어쓰기 요청을 입력해 주세요.")
            else:
                prog = st.progress(0); status = st.empty()
                status.write("25% · 선택 버전 세계관과 원고를 불러오는 중..."); prog.progress(25)
                result = ai_continue_from_selected_versions(selected_text, continue_req, int(continue_episode_no), continue_style)
                status.write("80% · 이어쓰기 결과를 새 버전으로 저장 중..."); prog.progress(80)
                p["latest_result"] = result
                p["latest_draft"] = result
                p["master_text"] = (p.get("master_text", "") + "\n\n" + result).strip()
                snap = auto_save_version(f"{int(continue_episode_no)}화 이어쓰기 · 선택버전", f"선택 버전 {selected_nos} 기준 이어쓰기 요청: {continue_req}", result)
                status.write(f"100% · 저장 완료: {snap.get('label')}"); prog.progress(100)
                st.success(f"선택 버전 기준 이어쓰기 완료 및 새 버전 저장: {snap.get('label')}")
                st.rerun()

        st.markdown("### ♻️ 선택 버전 복원")
        if len(selected_versions) != 1:
            st.info("복원은 한 번에 한 개 버전만 가능합니다. 하나만 체크해 주세요.")
        else:
            if st.button("♻️ 선택 버전으로 복원", key=f"restore_selected_version_v35_{version_no}"):
                p["master_text"] = primary_v.get("master_text") or primary_v.get("content", "")
                p["world_bible"] = primary_v.get("world_bible", "")
                p["characters"] = primary_v.get("characters", "")
                p["power_system"] = primary_v.get("power_system", "")
                p["episode_outline"] = primary_v.get("episode_outline", "")
                p["latest_draft"] = primary_v.get("latest_draft", "")
                p["latest_result"] = primary_v.get("latest_result") or primary_v.get("content", "")
                snap = auto_save_version("버전 복원", f"{primary_v.get('label')}에서 복원", export_current_text())
                st.success(f"복원 완료 및 새 버전 저장: {snap.get('label')}")
                st.rerun()

        st.markdown("### 🗑 선택 버전 삭제")
        st.markdown("<div class='danger-box'>삭제 후 복구가 어렵습니다. 중요한 버전은 먼저 다운로드하세요. V1/V2는 기본 보호됩니다.</div>", unsafe_allow_html=True)
        contains_protected = any(int(sv.get("version_no", 0)) <= 2 for sv in selected_versions)
        unlock = True
        if contains_protected:
            st.warning("선택 목록에 V1/V2 보호 버전이 포함되어 있습니다.")
            unlock = st.checkbox("V1/V2 보호 해제", key="unlock_delete_selected_v35")
        confirm_text = st.text_input("삭제 확인 문구 입력: DELETE", key="delete_confirm_selected_v35")
        if st.button("🗑 체크한 버전 삭제", key="delete_selected_versions_button_v35"):
            if contains_protected and not unlock:
                st.error("V1/V2 보호 해제가 필요합니다.")
            elif confirm_text.strip() != "DELETE":
                st.error("삭제하려면 DELETE 를 정확히 입력해 주세요.")
            else:
                deleted = []
                for no in selected_nos:
                    if delete_version_by_no(no):
                        deleted.append(no)
                st.success(f"삭제 완료: {', '.join(['V'+str(x) for x in deleted])}")
                st.rerun()

        if len(versions) >= 2:
            st.markdown("### 🔄 버전 비교")
            labels = [f"V{v.get('version_no')} · {v.get('action')} · {v.get('saved_at')}" for v in versions]
            ca, cb = st.columns(2)
            va_label = ca.selectbox("이전 버전", labels, index=min(1, len(labels)-1), key="compare_a_v35")
            vb_label = cb.selectbox("비교 버전", labels, index=0, key="compare_b_v35")
            if st.button("🔍 간단 비교"):
                va = versions[labels.index(va_label)]; vb = versions[labels.index(vb_label)]
                diff_text = f"""# 버전 비교

- 이전: {va_label}
- 비교: {vb_label}

## 글자수 변화
- {len(va.get('content','')):,}자 → {len(vb.get('content','')):,}자
- 차이: {len(vb.get('content','')) - len(va.get('content','')):+,}자

## 작업명
- 이전 작업: {va.get('action','')}
- 비교 작업: {vb.get('action','')}
"""
                st.markdown(f"<div class='ai-box'>{diff_text}</div>", unsafe_allow_html=True)

with tabs[4]:
    st.subheader("📥 다운로드")
    current_text = export_current_text()
    safe_base = slugify(f"{p.get('project_name','작품')}_V{p.get('current_version',0)}")
    c1, c2, c3 = st.columns(3)
    c1.download_button("📥 현재 작품 TXT", current_text, f"{safe_base}.txt", "text/plain")
    c2.download_button("📥 현재 작품 JSON", json.dumps(p, ensure_ascii=False, indent=2), f"{safe_base}.json", "application/json")
    docx = make_docx_bytes(p.get("project_name", "작품"), current_text)
    if docx:
        c3.download_button("📥 현재 작품 DOCX", docx, f"{safe_base}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        c3.info("DOCX는 python-docx 설치 시 사용 가능")
    zip_data = make_project_zip_bytes(p)
    st.download_button("📦 전체 프로젝트 ZIP 다운로드", zip_data, f"{safe_base}_backup.zip", "application/zip")
    st.markdown("### 백업 복원")
    restore = st.file_uploader("V37/V37/V32/V37 JSON 백업파일 복원", type=["json"], key="restore_v30")
    if restore is not None and st.button("♻️ 백업에서 복원"):
        try:
            data = json.loads(restore.read().decode("utf-8"))
            if "project_name" not in data:
                st.error("올바른 프로젝트 JSON이 아닙니다.")
            else:
                st.session_state.project_v30 = data
                save_project()
                st.success("복원 완료")
                st.rerun()
        except Exception as e:
            st.error(f"복원 실패: {e}")

with tabs[5]:
    st.subheader("📊 OpenAI 사용현황 V37")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("초기 충전", f"${OPENAI_INITIAL_CREDIT:.2f}")
    c2.metric("예상 사용", f"${st.session_state.openai_estimated_cost_v30:.4f}")
    c3.metric("예상/보정 잔액", f"${current_balance():.4f}")
    c4.metric("잔액 비율", f"{(current_balance()/OPENAI_INITIAL_CREDIT*100 if OPENAI_INITIAL_CREDIT else 0):.1f}%")
    st.progress(min(100, int((current_balance()/OPENAI_INITIAL_CREDIT*100) if OPENAI_INITIAL_CREDIT else 0)))
    st.link_button("💳 OpenAI Billing 실제 잔액 확인", "https://platform.openai.com/settings/organization/billing")
    actual = st.number_input("OpenAI Billing에서 확인한 실제 잔액($)", min_value=0.0, max_value=10000.0, value=float(current_balance()), step=0.01)
    if st.button("실제 잔액으로 보정"):
        st.session_state.openai_actual_balance_v30 = float(actual)
        st.success(f"앱 표시 잔액을 ${actual:.2f} 기준으로 보정했습니다.")
        st.rerun()
    st.dataframe(pd.DataFrame([
        {"항목": "Input Token", "값": st.session_state.openai_input_tokens_v30},
        {"항목": "Output Token", "값": st.session_state.openai_output_tokens_v30},
        {"항목": "Total Token", "값": st.session_state.openai_total_tokens_v30},
        {"항목": "AI 호출", "값": st.session_state.openai_call_count_v30},
        {"항목": "기본 모델", "값": OPENAI_MODEL},
    ]), use_container_width=True, hide_index=True)
    if st.button("사용량 기록 초기화"):
        for k in ["openai_input_tokens_v30", "openai_output_tokens_v30", "openai_total_tokens_v30", "openai_estimated_cost_v30", "openai_call_count_v30", "openai_actual_balance_v30"]:
            st.session_state[k] = DEFAULTS[k]
        st.rerun()



with tabs[6]:
    st.subheader("🎬 AI 스토리 분석실")
    st.markdown("""
<div class='card'>
<b>V37 신규 기능</b><br>
이전에 작성한 <b>구성안/시놉시스/원고</b>를 AI가 먼저 분석한 뒤,
스릴러·반전·코믹·감동·액션 등 원하는 장르를 선택하여 <b>이후 구성안</b>을 이어서 작성합니다.<br><br>
예: 기존 드라마 구성안을 업로드 → 스릴러 + 반전 + 코믹 선택 → 6~8화 구성안 자동 생성
</div>
""", unsafe_allow_html=True)

    story_source_mode = st.radio(
        "분석할 자료 선택",
        ["현재 작품", "작품보관함 버전", "파일 업로드 구성안/시놉"],
        horizontal=True,
        key="story_source_mode_v37"
    )

    story_source_text = ""
    selected_story_label = "현재 작품"

    if story_source_mode == "현재 작품":
        story_source_text = export_current_text()
        st.success(f"현재 작품 DB를 분석합니다. 글자수 {len(story_source_text):,}")

    elif story_source_mode == "작품보관함 버전":
        versions_for_story = p.get("versions", [])
        if not versions_for_story:
            st.warning("작품보관함에 저장된 버전이 없습니다.")
        else:
            story_labels = [f"V{v.get('version_no')} · {v.get('action')} · {v.get('saved_at')}" for v in versions_for_story]
            story_selected = st.selectbox("분석할 작품보관함 버전", story_labels, key="story_version_select_v37")
            sv = versions_for_story[story_labels.index(story_selected)]
            story_source_text = sv.get("content", "")
            selected_story_label = story_selected
            st.success(f"{story_selected} 기준으로 분석합니다. 글자수 {len(story_source_text):,}")

    else:
        story_upload = st.file_uploader("구성안/시놉시스/원고 업로드", type=["txt", "pdf", "docx"], key="story_upload_v37")
        if story_upload is not None:
            if st.button("📖 스토리 분석용 파일 읽기", key="read_story_upload_v37"):
                story_text = read_uploaded_file(story_upload)
                p["uploaded_name"] = story_upload.name
                p["uploaded_text"] = story_text
                save_project()
                st.success(f"업로드 저장 완료: {story_upload.name} / 글자수 {len(story_text):,}")
                st.rerun()
        if p.get("uploaded_text"):
            story_source_text = p.get("uploaded_text", "")
            selected_story_label = p.get("uploaded_name", "업로드 자료")
            st.success(f"현재 업로드 자료: {selected_story_label} / 글자수 {len(story_source_text):,}")

    st.markdown("### 장르/방향 선택")
    genre_choices = st.multiselect(
        "이후 구성안에 넣고 싶은 장르 요소",
        ["스릴러", "반전", "코믹", "감동", "액션", "미스터리", "공포", "로맨스", "정치암투", "복수극", "성장", "먼치킨", "휴먼드라마", "가족서사"],
        default=["스릴러", "반전"],
        key="story_genre_choices_v37"
    )

    sc1, sc2, sc3 = st.columns(3)
    output_type = sc1.selectbox(
        "결과물 형식",
        ["드라마 회차 구성안", "드라마 시놉시스", "웹소설 다음 회차 구성안", "영화 기획안", "OTT 시리즈 구성안", "웹툰 구성안", "자유 구성안"],
        index=0,
        key="story_output_type_v37"
    )
    continue_range = sc2.selectbox(
        "이어쓸 범위",
        ["다음 1화/1회", "다음 2~3화", "다음 4~6화", "다음 7~10화", "시즌 전체", "장편 전체 플롯"],
        index=1,
        key="story_continue_range_v37"
    )
    tone_level = sc3.selectbox(
        "강도/톤",
        ["균형형", "강한 반전", "스릴러 강하게", "코믹 많이", "감동 많이", "사이다 빠르게", "느린 빌드업", "대중 흥행형"],
        index=0,
        key="story_tone_level_v37"
    )

    story_request = st.text_area(
        "추가로 원하는 내용",
        height=130,
        placeholder="예: 초반 전개가 너무 빠르지 않게 하고, 주인공이 쉽게 강해지지 않게 해줘. 중간에 예상 못 한 반전과 코믹한 조력자를 넣어줘.",
        key="story_extra_request_v37"
    )

    if st.button("🎬 기존 구성안 분석 + 이후 구성안 생성 + 자동저장", type="primary", key="run_story_analysis_v37"):
        if not story_source_text.strip():
            st.error("분석할 구성안/시놉시스/원고가 없습니다. 현재 작품, 보관함 버전, 파일 업로드 중 하나를 선택해 주세요.")
        else:
            prog = st.progress(0); status = st.empty()
            status.write("20% · 기존 구성안/시놉시스 분석 중..."); prog.progress(20)
            result = ai_story_analysis_continue(
                story_source_text,
                genre_choices,
                output_type,
                continue_range,
                story_request,
                tone_level
            )
            status.write("80% · 분석 결과 및 이후 구성안 자동저장 중..."); prog.progress(80)
            p["latest_result"] = result
            p["latest_draft"] = result
            p["master_text"] = (p.get("master_text", "") + "\n\n" + result).strip() if p.get("master_text") else result
            detail = f"자료: {selected_story_label} / 장르: {', '.join(genre_choices)} / 형식: {output_type} / 범위: {continue_range}"
            snap = auto_save_version("스토리 분석실 · 이후 구성안 생성", detail, result)
            status.write(f"100% · 자동저장 완료: {snap.get('label')}"); prog.progress(100)
            st.success(f"스토리 분석 및 이후 구성안 생성 완료: {snap.get('label')}")
            st.rerun()

    if p.get("latest_result"):
        st.markdown("### 최근 스토리 분석/구성안 결과")
        st.markdown(f"<div class='ai-box'>{p.get('latest_result')[:10000]}</div>", unsafe_allow_html=True)



with tabs[7]:
    st.subheader("🤖 AI 공동작가 지시창")
    st.markdown("""
<div class='card'>
<b>여기는 작품을 직접 덮어쓰는 곳이 아니라, AI에게 작업 지시를 내리는 곳입니다.</b><br><br>
예: <b>1화 완성본 써줘</b>, <b>작품 DB 통합본으로 정리해줘</b>, <b>연나세계를 추가해줘</b>, <b>주인공을 더 먼치킨으로 강화해줘</b><br><br>
실행 결과는 기존 V1/V2를 삭제하지 않고 <b>항상 새 버전</b>으로 자동 저장됩니다.
</div>
""", unsafe_allow_html=True)

    if "ai_command_text_v31" not in st.session_state:
        st.session_state.ai_command_text_v31 = ""

    st.markdown("### 빠른 작업 버튼")
    b1, b2, b3 = st.columns(3)
    if b1.button("📖 1화 완성본 써줘"):
        st.session_state.ai_command_text_v31 = "현재 작품 DB와 마스터 원고를 바탕으로 1화 완성본을 웹소설 형식으로 작성해줘. 주인공의 첫 각성, 강한 후킹, 다음 화가 궁금해지는 엔딩을 넣어줘."
    if b2.button("📚 작품 DB 통합본 정리"):
        st.session_state.ai_command_text_v31 = "현재 흩어진 설정을 하나의 작품 DB 통합본으로 정리해줘. 세계관, 등장인물, 세력, 무공/능력, 회차 방향, 장기 떡밥을 보기 쉽게 정리해줘."
    if b3.button("⚔️ 세력/무공 정리"):
        st.session_state.ai_command_text_v31 = "현재 작품에 맞게 세력, 무공, 헌터 능력, 신격, 판타지 마법 체계를 충돌 없이 정리하고 더 멋지게 확장해줘."

    b4, b5, b6 = st.columns(3)
    if b4.button("📖 2~5화 연속 플롯"):
        st.session_state.ai_command_text_v31 = "현재 1화 이후 2화부터 5화까지의 연속 플롯을 작성해줘. 각 화마다 핵심 사건, 전투, 감정선, 반전, 엔딩 후킹을 넣어줘."
    if b5.button("🌌 연나세계 추가"):
        st.session_state.ai_command_text_v31 = "작품에 연나세계를 추가해줘. 기존 무림세계, 현대 헌터세계, 판타지세계와 충돌하지 않게 연결하고, 독자가 예상 못 할 반전과 감동 요소를 넣어줘."
    if b6.button("📈 100화 장기 플롯"):
        st.session_state.ai_command_text_v31 = "현재 작품을 100화 이상 장기 연재할 수 있도록 큰 장기 플롯을 설계해줘. 1부, 2부, 3부, 최종부로 나누고 각 부의 핵심 적, 성장, 반전을 정리해줘."

    mode = st.selectbox("작업 모드", ["자유 지시", "1화/본문 작성", "작품 DB 정리", "세계관 확장", "등장인물 강화", "세력/무공 정리", "장기 플롯 설계"], index=0)
    command = st.text_area(
        "AI 공동작가에게 지시",
        key="ai_command_text_v31",
        height=180,
        placeholder="예: 현재 V2 내용은 유지하고, 주인공이 주신의 의뢰로 연나세계에 들어가는 내용을 추가해줘. 전투는 더 압도적으로, 감동과 반전도 넣어줘."
    )

    if st.button("🤖 AI 공동작가 실행 + 새 버전 자동저장", type="primary"):
        if not command.strip():
            st.error("AI에게 지시할 내용을 입력해 주세요.")
        else:
            prog = st.progress(0); status = st.empty()
            status.write("20% · 기존 V1/V2 및 현재 작품 DB 불러오는 중..."); prog.progress(20)
            result = ai_cowriter_command(command, mode)
            status.write("75% · 기존 내용 보존 후 새 작업 결과 자동저장 중..."); prog.progress(75)
            p = st.session_state.project_v30

            # 기존 마스터 원고는 삭제하지 않고, 새 AI 결과를 아래에 누적합니다.
            preserved = p.get("master_text", "").strip()
            section_title = f"\n\n\n---\n# V37 AI 공동작가 작업 결과 · {now_str()}\n\n"
            p["master_text"] = (preserved + section_title + result).strip() if preserved else result
            p["latest_result"] = result
            p["latest_draft"] = result
            snap = auto_save_version(f"AI 공동작가 지시 · {mode}", command, export_current_text())
            status.write(f"100% · 새 버전 자동저장 완료: {snap.get('label')}"); prog.progress(100)
            st.success(f"AI 작업 완료. 기존 V1/V2는 보존하고 새 버전으로 저장했습니다: {snap.get('label')}")
            st.rerun()

    if p.get("latest_result"):
        st.markdown("### 최근 AI 공동작가 결과")
        st.markdown(f"<div class='ai-box'>{p.get('latest_result')[:10000]}</div>", unsafe_allow_html=True)



with tabs[8]:
    st.subheader("⚙️ V37 사용 가이드")
    st.markdown("""
<div class='card'>
<b>V37 핵심 사용법</b><br><br>
1. <b>작품 생성</b>에서 한 문장 아이디어를 입력합니다.<br>
2. AI가 작품을 만들면 자동으로 <b>작품 보관함 V1</b>에 저장됩니다.<br>
3. <b>🤖 AI 공동작가</b> 탭에 “1화 완성본 써줘”, “작품 DB 통합본으로 정리해줘”처럼 지시하면 자동으로 <b>새 버전</b>이 저장됩니다.<br>
4. <b>🎬 스토리 분석실</b>에서 기존 구성안/시놉을 분석하고, 스릴러·반전·코믹 등 장르를 선택해 이후 구성안을 생성합니다.<br>5. <b>작품 보관함</b>에서 기준이 될 버전을 체크한 뒤 <b>선택 버전 기준 이어쓰기</b>를 실행합니다.<br>
6. <b>작품 보관함</b>에서 이전 버전을 확인하고 복원할 수 있습니다.<br>
7. <b>작품 보관함</b>에서 선택 버전을 편집/복원/삭제/다운로드할 수 있습니다.<br>
8. <b>다운로드</b>에서 TXT/DOCX/JSON/ZIP으로 백업하세요.<br><br>
<b>중요</b><br>
AI 작업 결과는 성공 시 자동저장됩니다. 기존 V1/V2는 삭제하지 않고, 항상 새 버전으로 작품 보관함에 보존됩니다. V37에서는 작품보관함에서 선택 버전을 편집/다운로드/복원/삭제할 수 있습니다.
</div>
""", unsafe_allow_html=True)
